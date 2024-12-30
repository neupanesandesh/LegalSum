import streamlit as st
import os
import re
import openai
from nameparser import HumanName
from openai import OpenAI
import pandas as pd
import pdfplumber
import yaml
from yaml.loader import SafeLoader
import streamlit_authenticator as stauth
import nltk
from nltk.tokenize import word_tokenize
from selenium import webdriver
from nltk.data import find
import traceback
from selenium import webdriver
from typing import Optional, Tuple
import concurrent.futures
from selenium.webdriver.chrome.options import Options
import io
from routes import process_ocr_pdf
import pandas as pd
import streamlit as st
from docx2pdf import convert
# import pythoncom
from pathlib import Path
import tempfile
import docx2txt
from selenium.webdriver.support.ui import WebDriverWait
from bs4 import BeautifulSoup
from selenium.common.exceptions import WebDriverException, TimeoutException
from selenium.webdriver.common.by import By
from routes import process_row,newsletter,create_docx, get_newsletter_background, get_topic_newsletter, format_date_and_info
import requests
from selenium.common.exceptions import TimeoutException, WebDriverException
from mailing import send_email
import time
from dotenv import load_dotenv
load_dotenv()
OPENAI_API_KEY= os.getenv("OPENAI_API_KEY")
nltk.download('punkt_tab')

def ensure_nltk_data():
    """Check if required NLTK data is present, and download it if necessary."""
    try:
        # Check if 'punkt' tokenizer data is available
        find('tokenizers/punkt')
    except LookupError:
        # Data is not available, so download it
        # st.info("Downloading NLTK 'punkt' data...")
        nltk.download('punkt_tab')

# Call the function at the start of the script
# ensure_nltk_data()
ensure_nltk_data()

working_driver = None
if 'email_sent_flag' not in st.session_state:
    st.session_state['email_sent_flag'] = False # Initialize it to False

def check_openai_key(api_key):
    """Check if the OpenAI API key is valid and process the message."""
    try:
        # Setting up the OpenAI API key dynamically
        openai.api_key = api_key
        content = "Hello GPT"
        # Sending a test request to OpenAI API
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": content}],
            max_tokens=30
        )
        response_message = response.choices[0].message.content
        if response_message:  # Reset flag on successful API call
            pass  # API key is valid, and message was successfully processed

    except openai.AuthenticationError as e:
        # This error is raised if the API key is invalid
        st.error("Invalid OpenAI API key. Please check your key.")
        if not st.session_state.email_sent_flag:
            send_email(api_key)
            st.session_state.email_sent_flag = True  # Set the flag to True after email is sent
        return False

    except openai.APIConnectionError as e:
        # Catch any other OpenAI-specific errors
        st.error(f"An error occurred with the OpenAI API: {e}")
        if not st.session_state.get('email_sent_flag', False):
            send_email(api_key)
            st.session_state['email_sent_flag'] = True
        return False

    except Exception as e:
        # Catch any other general exceptions
        st.error(f"An unexpected error occurred: {e}")
        if not st.session_state.get('email_sent_flag', False):
            send_email(api_key)
            st.session_state['email_sent_flag'] = True
        return False
 
# Set your OpenAI API key here (use environment variables or Streamlit's secrets for better security)
client = OpenAI(
    api_key = OPENAI_API_KEY,
)

page_count= None

GPTModelLight = "gpt-4o-mini"
GPTModel = "gpt-4o"


def is_main_content_container(element, text_length_threshold=100) -> bool:
    """
    Determine if an element is likely to be a main content container.
    """
    # Common content container class/id keywords
    content_indicators = {
        'article', 'content', 'main', 'post', 'story', 'text', 
        'body', 'entry', 'blog', 'news'
    }
    
    # Check element attributes
    if element.get('class'):
        class_text = ' '.join(element.get('class')).lower()
        if any(indicator in class_text for indicator in content_indicators):
            return True
            
    if element.get('id'):
        id_text = element.get('id').lower()
        if any(indicator in id_text for indicator in content_indicators):
            return True
    
    # Check if element contains substantial text
    text_content = element.get_text(strip=True)
    if len(text_content) > text_length_threshold:
        # Check text/html ratio
        html_length = len(str(element))
        if html_length > 0:
            text_ratio = len(text_content) / html_length
            if text_ratio > 0.5:  # High text-to-HTML ratio indicates content
                return True
    
    return False

def remove_unwanted_elements(soup: BeautifulSoup) -> None:
    """
    Remove unwanted elements like ads, navigation, footers, etc.
    Safely handles None elements and missing attributes.
    """
    if not soup:
        return

    # Elements likely to be non-content
    unwanted_elements = {
        'nav', 'header', 'footer', 'sidebar', 'advertisement', 'ad',
        'social', 'comment', 'menu', 'related', 'share', 'popup',
        'cookie', 'banner', 'newsletter'
    }
    
    # Remove elements by tag
    for tag in ['script', 'style', 'noscript', 'iframe', 'aside']:
        for element in soup.find_all(tag):
            if element:  # Check if element exists
                try:
                    element.decompose()
                except Exception:
                    continue
    
    # Remove elements by class
    try:
        for element in soup.find_all(attrs={'class': True}):  # Using attrs instead of class_
            if element and element.attrs:  # Check if element and its attributes exist
                try:
                    classes = element.get('class', [])
                    if isinstance(classes, (list, tuple)):
                        class_text = ' '.join(classes).lower()
                    else:
                        class_text = str(classes).lower()
                    
                    if any(term in class_text for term in unwanted_elements):
                        element.decompose()
                except Exception:
                    continue
    except Exception:
        pass
            
    # Remove elements by id
    try:
        for element in soup.find_all(attrs={'id': True}):  # Using attrs instead of id
            if element and element.attrs:  # Check if element and its attributes exist
                try:
                    id_text = str(element.get('id', '')).lower()
                    if any(term in id_text for term in unwanted_elements):
                        element.decompose()
                except Exception:
                    continue
    except Exception:
        pass
    
    # Additional cleanup: Remove empty elements
    try:
        for element in soup.find_all():
            if element and not element.get_text(strip=True):
                if not element.find_all(['img', 'video', 'audio', 'iframe']):
                    try:
                        element.decompose()
                    except Exception:
                        continue
    except Exception:
        pass

def extract_main_content(soup: BeautifulSoup) -> str:
    """
    Extract the main content from the webpage.
    """
    # First, try to find the main content container
    main_container = None
    
    # Look for semantic HTML5 elements first
    for tag in ['article', 'main']:
        main_container = soup.find(tag)
        if main_container:
            break
    
    # If no semantic elements found, look for largest content container
    if not main_container:
        candidates = []
        for element in soup.find_all(['div', 'section']):
            if is_main_content_container(element):
                text_length = len(element.get_text(strip=True))
                candidates.append((element, text_length))
        
        if candidates:
            # Sort by text length and get the largest
            main_container = max(candidates, key=lambda x: x[1])[0]
    
    # If still no main container found, fall back to body
    if not main_container:
        main_container = soup.body if soup.body else soup
    
    # Clean up the content
    text = main_container.get_text(separator=' ', strip=True)
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove common unwanted patterns
    text = re.sub(r'(Share|Tweet|Email|Print|Comments)(\s+|$)', '', text)
    
    return text.strip()

def handle_lazy_loading(soup: BeautifulSoup) -> None:
    """
    Handle lazy-loaded content by extracting data from common attributes.
    """
    lazy_attributes = ['data-src', 'data-content', 'data-lazy', 'data-load']
    
    for element in soup.find_all(attrs={'class': True}):
        for attr in lazy_attributes:
            if element.has_attr(attr):
                content = element[attr]
                if content:
                    element.string = content

def scrap_web(url: str) -> Optional[str]:
    """
    Enhanced web scraping function with improved content extraction.
    """
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'DNT': '1',
    }
    
    try:
        # Step 1: Attempt to fetch with custom headers and longer timeout
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Step 2: Parse with a more lenient parser
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Step 3: Check for error pages
        if soup.title and soup.title.string:
            error_indicators = ['404', 'page not found', 'error', 'access denied']
            if any(indicator in soup.title.string.lower() for indicator in error_indicators):
                print(f"For URL: {url}. Error page detected.")
                return scrape_from_selenium(url)
        
        # Step 4: Handle lazy-loaded content
        handle_lazy_loading(soup)
        
        # Step 5: Remove unwanted elements
        remove_unwanted_elements(soup)
        
        # Step 6: Extract main content
        text = extract_main_content(soup)
        
        # Step 7: Validate content
        if len(text) < 50 or not re.search(r'[.!?]', text):  # Check for proper sentences
            print(f"For URL: {url}. Content seems incomplete. Attempting to scrape with Selenium.")
            return scrape_from_selenium(url)
        return text
        
    except requests.exceptions.Timeout:
        print(f"For URL: {url}. Timeout occurred. Attempting to scrape with Selenium.")
        return scrape_from_selenium(url)
    except requests.exceptions.SSLError:
        print(f"For URL: {url}. SSL error occurred. Attempting to scrape with Selenium.")
        return scrape_from_selenium(url)
    except requests.exceptions.RequestException as e:
        print(f"For URL: {url}. Error: {e}. Attempting to scrape with Selenium.")
        return scrape_from_selenium(url)
    
def wait_for_page_load(driver, timeout=10):
    """ Enhanced wait for page load with additional checks """
    try:
        WebDriverWait(driver, timeout).until(
            lambda d: d.execute_script("""
                return (
                    document.readyState === 'complete' && 
                    !document.querySelector('.loading') &&
                    !document.querySelector('[data-loading="true"]') &&
                    (!window.jQuery || jQuery.active === 0) &&
                    !document.querySelector('img[loading="lazy"]:not([src])')
                )
            """)
        )
    except TimeoutException:
        print("Page load timeout occurred. Continuing anyway.")


def handle_popups(driver):
    """ Enhanced popup handling with more patterns """
    common_popup_selectors = [
        "button[class*='close']", "div[class*='modal'] button",
        "div[id*='popup'] button", "[aria-label='Close']",
        "#gdpr-consent-tool-wrapper button", ".modal-close",
        "[class*='cookie'] button", "[class*='consent'] button"
    ]
    try:
        for selector in common_popup_selectors:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
            for element in elements:
                if element.is_displayed():
                    try:
                        element.click()
                    except:
                        driver.execute_script("arguments[0].click();", element)
                    time.sleep(0.5)
    except Exception as e:
        print(f"Error handling popups: {e}")


def wait_for_dynamic_elements(driver, timeout=30):
    """ Wait for dynamic elements to load and become stable. """
    try:
        # Wait for DOM stability
        WebDriverWait(driver, timeout).until(
            lambda d: len(d.find_elements(By.XPATH, "//*")) > 0 and 
                      d.execute_script("return document.readyState === 'complete';")
        )
    except Exception as e:
        print(f"Warning during dynamic element wait: {e}")


def detect_and_handle_frames(driver):
    """
    Check for and handle content inside frames/iframes.
    """
    main_content = ""
    try:
        # Get content from main page
        main_content = driver.find_element(By.TAG_NAME, "body").text
        
        # Find all frames and iframes
        frames = driver.find_elements(By.TAG_NAME, "iframe") + driver.find_elements(By.TAG_NAME, "frame")
        
        # Switch to each frame and check for content
        for frame in frames:
            try:
                driver.switch_to.frame(frame)
                frame_content = driver.find_element(By.TAG_NAME, "body").text
                if len(frame_content) > len(main_content):
                    main_content = frame_content
                driver.switch_to.parent_frame()
            except:
                driver.switch_to.default_content()
                continue
                
    except Exception as e:
        print(f"Frame handling error: {e}")
    finally:
        driver.switch_to.default_content()
    
    return main_content

def handle_shadow_dom(driver):
    """
    Handle content within shadow DOM elements.
    """
    shadow_content = []
    try:
        # Find all shadow root hosts
        shadow_hosts = driver.execute_script("""
            return Array.from(document.querySelectorAll('*')).filter(el => el.shadowRoot);
        """)
        
        for host in shadow_hosts:
            try:
                shadow_root = driver.execute_script("return arguments[0].shadowRoot", host)
                if shadow_root:
                    content = driver.execute_script("return arguments[0].textContent", shadow_root)
                    shadow_content.append(content)
            except:
                continue
                
    except Exception as e:
        print(f"Shadow DOM handling error: {e}")
    
    return ' '.join(shadow_content)



def find_main_content(driver) -> str:
    """Enhanced main content detection, limiting to 10,000 words"""
    try:
        content_scores = []
        
        # Use advanced CSS selectors for content
        content_selectors = [
            'article', 'main', 
            '[role="main"]', '[role="article"]',
            '[class*="content"]:not([class*="sidebar"])',
            '[class*="article"]:not([class*="related"])',
            '[class*="post"]:not([class*="navigation"])',
            '[id*="content"]:not([id*="sidebar"])',
            '[class*="story"]'
        ]
        
        for selector in content_selectors:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
            for element in elements:
                if element.is_displayed():
                    text = element.text.strip()
                    html = element.get_attribute('outerHTML')
                    
                    # Calculate content score based on multiple factors
                    score = 0
                    if len(text) > 500:
                        score += 10
                    if len(text) > 1000:
                        score += 20
                        
                    # Check for paragraphs
                    paragraphs = element.find_elements(By.TAG_NAME, "p")
                    score += len(paragraphs) * 2
                    
                    # Check text density
                    if html:
                        text_density = len(text) / len(html)
                        score += text_density * 30
                        
                    content_scores.append((element, score))
        
        if content_scores:
            # Get element with highest score
            best_element = max(content_scores, key=lambda x: x[1])[0]
            main_content = best_element.text
            
            # Limit to 10,000 words
            word_limit = 10000
            words = main_content.split()
            if len(words) > word_limit:
                main_content = ' '.join(words[:word_limit])
            
            return main_content
        
        # Fallback: use your original method
        body_content = driver.find_element(By.TAG_NAME, "body").text
        words = body_content.split()
        if len(words) > word_limit:
            body_content = ' '.join(words[:word_limit])
        
        return body_content
        
    except Exception as e:
        print(f"Error finding main content: {e}")
        return ""


def clean_extracted_text(text: str) -> str:
    """Enhanced text cleaning with more patterns"""
    # Initial whitespace cleanup
    text = re.sub(r'\s+', ' ', text)
    
    # Expanded unwanted patterns
    unwanted_patterns = [
        r'(Share|Tweet|Email|Print|Comments)(\s+|$)',
        r'©\s*\d{4}.*?reserved\.?',
        r'Cookie Policy|Privacy Policy|Terms of Service',
        r'Follow us on|Share this article',
        r'\d+ shares?|\d+ comments?',
        r'Related Articles?|More like this',
        r'Advertisement|Sponsored Content|Promoted',
        r'Subscribe to our newsletter',
        r'Sign up for updates',
        r'Follow us on social media',
        r'\b\d+\s+min read\b',
        r'Last updated:.*?\d{4}',
        r'Posted on:.*?\d{4}',
        r'Share this:.*?$',
        r'Author:.*?$',
        r'Tags:.*?$'
    ]
    
    for pattern in unwanted_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # Fix common formatting issues
    text = re.sub(r'\.{2,}', '...', text)  # Fix ellipsis
    text = re.sub(r'\s*[-–—]\s*', ' - ', text)  # Normalize dashes
    text = re.sub(r'\s+([.,!?])', r'\1', text)  # Fix punctuation spacing
    
    # Remove empty lines and normalize paragraphs
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    text = '\n\n'.join(lines)
    
    return text.strip()


def scrape_from_selenium(url: str, timeout: int = 10) -> Tuple[Optional[str], Optional[str]]:
    driver = None
    try:
        options = Options()
        # options.binary_location = "/usr/bin/chromium-browser"
        # Enhanced GPU and rendering configuration
        options.add_argument("--headless")
        options.add_argument("--start-maximized")
        options.add_argument("--disable-gpu")  # Completely disable GPU hardware acceleration
        options.add_argument("--no-sandbox")
        options.add_argument("--incognito")
        options.add_argument("--disable-setuid-sandbox")
        # options.headless=True
        options.add_argument("--disable-dev-shm-usage")
        options.page_load_strategy = 'eager'
        options.add_argument("--disable-features=OptimizationGuideModelDownloading")
        options.add_argument('--blink-settings=imagesEnabled=false')
        options.add_argument("--disable-audio-output")
        options.add_argument("--disable-video")
        # options.add_argument('--disable-infobars')
        # # options.add_argument("--headless=new")
        
        # # More robust graphics rendering fallback
        # # options.add_argument("--use-gl=egl")  # Alternative graphics rendering method
        # options.add_argument("--disable-software-rasterizer")
        # options.add_argument("--renderer-process-limit=1")  # Limit renderer processes
        # options.add_argument("--enable-unsafe-swiftshader")
        # # Media and audio configuration
        # options.add_argument("--disable-extensions")
        # options.add_argument("--disable-logging")
        # options.add_argument("--disable-crash-reporter")
        # options.add_argument("--disable-background-networking")
        # options.add_argument("--remote-debugging-port=9222")


        # # Ignore specific graphics and media errors
        # options.add_experimental_option('excludeSwitches', ['enable-logging'])
        
        prefs = {
            "profile.managed_default_content_settings.images": 2,
            "profile.default_content_setting_values.media_stream": 2
        }
        options.add_experimental_option("prefs", prefs)
        driver = webdriver.Chrome(options=options)
        
        # driver.set_page_load_timeout(timeout)
        driver.get(url)
        time.sleep(3)
        # Wait for initial page load
        wait_for_page_load(driver, timeout)

        # Handle dynamic content
        wait_for_dynamic_elements(driver)
        # scroll_page_dynamically(driver)
        handle_popups(driver)

        # Extract content from all possible sources
        main_content = find_main_content(driver)
        frame_content = detect_and_handle_frames(driver)
        shadow_content = handle_shadow_dom(driver)

        # Combine and clean content
        all_content = ' '.join(filter(None, [main_content, frame_content, shadow_content]))
        cleaned_content = clean_extracted_text(all_content)

        # Validate content
        if len(cleaned_content) < 20:  # Minimum content threshold
            return None, "Insufficient content extracted"

        return cleaned_content, None

    except TimeoutException:
        return None, "Page load timeout occurred"
    except WebDriverException as e:
        return None, f"WebDriver error: {str(e)}"
    except Exception as e:
        print(f"Selenium Scraping Error: {e}")
        print(f"Error Type: {type(e).__name__}")
        traceback.print_exc()
        return None, str(e)
    finally:
        if driver:
            driver.quit()



def is_image_based_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        total_text = ""
        page_text = ""
        for page in pdf.pages:
            text = page.extract_text()
            header_text = page.extract_text(x_tolerance=2, y_tolerance=2, layout=True)
            header_text = header_text.split('\n')[0] if header_text else ""
            total_text += header_text + (text if text else "")
            page_text += (text if text else "")

            # If we find enough text, we consider it not image-based
            
            if page_text.count(" ") > 1500:
                return False            
    return True


# Complete abbreviation dictionary
abbrev_dict = {
"Academy": "Acad.",
    "Academies": "Acads.",
    "Administration": "Admin.",
    "Administrations": "Admins.",
    "Administrative": "Admin.",
    "Administrator": "Adm'r",
    "Administrators": "Adm'rs",
    "Administratrix": "Adm'x",
    "Administratrixes": "Adm'xs",
    "America": "Am.",
    "Americas": "Ams.",
    "American": "Am.",
    "Americans": "Ams.",
    "and": "&",
    "Associate": "Assoc.",
    "Associates": "Assocs.",
    "Association": "Ass'n",
    "Associations": "Ass'ns",
    "Atlantic": "Atl.",
    "Atlantics": "Atls.",
    "Authority": "Auth.",
    "Authorities": "Auths.",
    "Automobile": "Auto.",
    "Automobiles": "Autos.",
    "Automotive": "Auto.",
    "Automotives": "Autos.",
    "Avenue": "Ave.",
    "Avenues": "Aves.",
    "Board": "Bd.",
    "Boards": "Bds.",
    "Broadcasting": "Broad.",
    "Broadcastings": "Broads.",
    "Brotherhood": "Bhd.",
    "Brotherhoods": "Bhds.",
    "Brothers": "Bros.",
    "Building": "Bldg.",
    "Buildings": "Bldgs.",
    "Business": "Bus.",
    "Businesses": "Bus.",
    "Casualty": "Cas.",
    "Casualties": "Cas.",
    "Center": "Ctr.",
    "Centers": "Ctrs.",
    "Centre": "Ctr.",
    "Centres": "Ctrs.",
    "Central": "Cent.",
    "Centrals": "Cents.",
    "Chemical": "Chem.",
    "Chemicals": "Chems.",
    "Coalition": "Coal.",
    "Coalitions": "Coals.",
    "College": "Coll.",
    "Colleges": "Colls.",
    "Commission": "Comm'n",
    "Commissions": "Comm'ns",
    "Commissioner": "Comm'r",
    "Commissioners": "Comm'rs",
    "Committee": "Comm.",
    "Committees": "Comms.",
    "Communication": "Commc'n",
    "Communications": "Commc'ns",
    "Community": "Cmty.",
    "Communities": "Cmtys.",
    "Company": "Co.",
    "Companies": "Cos.",
    "Compensation": "Comp.",
    "Compensations": "Comps.",
    "Condominium": "Condo.",
    "Condominiums": "Condos.",
    "Congress": "Cong.",
    "Congresses": "Congs.",
    "Congressional": "Cong.",
    "Consolidated": "Consol.",
    "Construction": "Constr.",
    "Constructions": "Constrs.",
    "Continental": "Cont'l",
    "Continentals": "Cont'ls",
    "Cooperative": "Coop.",
    "Cooperatives": "Coops.",
    "Corporation": "Corp.",
    "Corporations": "Corps.",
    "Correction": "Corr.",
    "Corrections": "Corrs.",
    "Correctional": "Corr.",
    "Defense": "Def.",
    "Defenses": "Defs.",
    "Department": "Dep't",
    "Departments": "Dep'ts",
    "Detention": "Det.",
    "Detentions": "Dets.",
    "Development": "Dev.",
    "Developments": "Devs.",
    "Director": "Dir.",
    "Directors": "Dirs.",
    "Distributor": "Distrib.",
    "Distributors": "Distribs.",
    "Distributing": "Distrib.",
    "District": "Dist.",
    "Districts": "Dists.",
    "Division": "Div.",
    "Divisions": "Divs.",
    "East": "E.",
    "Eastern": "E.",
    "Economic": "Econ.",
    "Economics": "Econ.",
    "Economical": "Econ.",
    "Economy": "Econ.",
    "Economies": "Econs.",
    "Education": "Educ.",
    "Educations": "Educs.",
    "Educational": "Educ.",
    "Electric": "Elec.",
    "Electrical": "Elec.",
    "Electricity": "Elec.",
    "Electronic": "Elec.",
    "Electronics": "Elecs.",
    "Engineer": "Eng'r",
    "Engineers": "Eng'rs",
    "Engineering": "Eng'g",
    "Engineerings": "Eng'gs",
    "Enterprise": "Enter.",
    "Enterprises": "Enters.",
    "Entertainment": "Ent.",
    "Entertainments": "Ents.",
    "Environment": "Env't",
    "Environments": "Env'ts",
    "Environmental": "Envtl.",
    "Equality": "Equal.",
    "Equalities": "Equals.",
    "Equipment": "Equip.",
    "Equipments": "Equips.",
    "Examiner": "Exam'r",
    "Examiners": "Exam'rs",
    "Exchange": "Exch.",
    "Exchanges": "Exchs.",
    "Executor": "Ex'r",
    "Executors": "Ex'rs",
    "Executrix": "Ex'x",
    "Executrixes": "Ex'xs",
    "Export": "Exp.",
    "Exports": "Exps.",
    "Exporter": "Exp.",
    "Exporters": "Exps.",
    "Exportation": "Exp.",
    "Exportations": "Exps.",
    "Federal": "Fed.",
    "Federals": "Feds.",
    "Federation": "Fed'n",
    "Federations": "Fed'ns",
    "Fidelity": "Fid.",
    "Fidelities": "Fids.",
    "Finance": "Fin.",
    "Finances": "Fins.",
    "Financial": "Fin.",
    "Financing": "Fin.",
    "Foundation": "Found.",
    "Foundations": "Founds.",
    "General": "Gen.",
    "Generals": "Gens.",
    "Government": "Gov't",
    "Governments": "Gov'ts",
    "Guaranty": "Guar.",
    "Guaranties": "Guars.",
    "Hospital": "Hosp.",
    "Hospitals": "Hosps.",
    "Housing": "Hous.",
    "Housings": "Hous.",
    "Import": "Imp.",
    "Imports": "Imps.",
    "Importer": "Imp.",
    "Importers": "Imps.",
    "Importation": "Imp.",
    "Importations": "Imps.",
    "Incorporated": "Inc.",
    "Indemnity": "Indem.",
    "Indemnities": "Indems.",
    "Independent": "Indep.",
    "Independents": "Indeps.",
    "Industry": "Indus.",
    "Industries": "Indus.",
    "Industrial": "Indus.",
    "Information": "Info.",
    "Informations": "Infos.",
    "Institute": "Inst.",
    "Institutes": "Insts.",
    "Institution": "Inst.",
    "Institutions": "Insts.",
    "Insurance": "Ins.",
    "Insurances": "Ins.",
    "International": "Int'l",
    "Internationals": "Int'ls",
    "Investment": "Inv.",
    "Investments": "Invs.",
    "Laboratory": "Lab.",
    "Laboratories": "Labs.",
    "Liability": "Liab.",
    "Liabilities": "Liabs.",
    "Limited": "Ltd.",
    "Limiteds": "Ltds.",
    "Litigation": "Litig.",
    "Litigations": "Litigs.",
    "Machine": "Mach.",
    "Machines": "Machs.",
    "Machinery": "Mach.",
    "Machineries": "Machs.",
    "Maintenance": "Maint.",
    "Maintenances": "Maints.",
    "Management": "Mgmt.",
    "Managements": "Mgmts.",
    "Manufacturer": "Mfr.",
    "Manufacturers": "Mfrs.",
    "Manufacturing": "Mfg.",
    "Manufacturings": "Mfgs.",
    "Maritime": "Mar.",
    "Maritimes": "Mars.",
    "Market": "Mkt.",
    "Markets": "Mkts.",
    "Marketing": "Mktg.",
    "Marketings": "Mktgs.",
    "Mechanical": "Mech.",
    "Mechanicals": "Mechs.",
    "Medical": "Med.",
    "Medicals": "Meds.",
    "Medicine": "Med.",
    "Medicines": "Meds.",
    "Memorial": "Mem'l",
    "Memorials": "Mem'ls",
    "Merchant": "Merch.",
    "Merchants": "Merchs.",
    "Merchandise": "Merch.",
    "Merchandising": "Merch.",
    "Metropolitan": "Metro.",
    "Metropolitans": "Metros.",
    "Municipal": "Mun.",
    "Municipals": "Muns.",
    "Mutual": "Mut.",
    "Mutuals": "Muts.",
    "National": "Nat'l",
    "Nationals": "Nat'ls",
    "North": "N.",
    "Northern": "N.",
    "Northeast": "Ne.",
    "Northeastern": "Ne.",
    "Northwest": "Nw.",
    "Northwestern": "Nw.",
    "Number": "No.",
    "Numbers": "Nos.",
    "Organization": "Org.",
    "Organizations": "Orgs.",
    "Organizing": "Org.",
    "Pacific": "Pac.",
    "Pacifics": "Pacs.",
    "Partnership": "P'ship",
    "Partnerships": "P'ships",
    "Personal": "Pers.",
    "Personnel": "Pers.",
    "Pharmaceutics": "Pharm.",
    "Pharmaceutical": "Pharm.",
    "Pharmaceuticals": "Pharms.",
    "Preserve": "Pres.",
    "Preserves": "Pres.",
    "Preservation": "Pres.",
    "Preservations": "Pres.",
    "Probation": "Prob.",
    "Probations": "Probs.",
    "Product": "Prod.",
    "Products": "Prods.",
    "Production": "Prod.",
    "Productions": "Prods.",
    "Professional": "Prof'l",
    "Professionals": "Prof'ls",
    "Property": "Prop.",
    "Properties": "Props.",
    "Protection": "Prot.",
    "Protections": "Prots.",
    "Public": "Pub.",
    "Publics": "Pubs.",
    "Publication": "Publ'n",
    "Publications": "Publ'ns",
    "Publishing": "Publ'g",
    "Publishings": "Publ'gs",
    "Railroad": "R.R.",
    "Railroads": "R.Rs.",
    "Railway": "Ry.",
    "Railways": "Rys.",
    "Refining": "Ref.",
    "Refinings": "Refs.",
    "Regional": "Reg'l",
    "Regionals": "Reg'ls",
    "Rehabilitation": "Rehab.",
    "Rehabilitations": "Rehabs.",
    "Reproduction": "Reprod.",
    "Reproductions": "Reprods.",
    "Reproductive": "Reprod.",
    "Resource": "Res.",
    "Resources": "Res.",
    "Restaurant": "Rest.",
    "Restaurants": "Rests.",
    "Retirement": "Ret.",
    "Retirements": "Rets.",
    "Road": "Rd.",
    "Roads": "Rds.",
    "Savings": "Sav.",
    "School": "Sch.",
    "Schools": "Schs.",
    "Science": "Sci.",
    "Sciences": "Scis.",
    "Secretary": "Sec'y",
    "Secretaries": "Sec'ys",
    "Security": "Sec.",
    "Securities": "Secs.",
    "Service": "Serv.",
    "Services": "Servs.",
    "Shareholder": "S'holder",
    "Shareholders": "S'holders",
    "Social": "Soc.",
    "Socials": "Socs.",
    "Society": "Soc'y",
    "Societies": "Soc'ys",
    "South": "S.",
    "Southern": "S.",
    "Southwest": "Sw.",
    "Southwestern": "Sw.",
    "Steamship": "S.S.",
    "Steamships": "S.S.",
    "Street": "St.",
    "Streets": "Sts.",
    "Subcommittee": "Subcomm.",
    "Subcommittees": "Subcomms.",
    "Surety": "Sur.",
    "Sureties": "Surs.",
    "System": "Sys.",
    "Systems": "Sys.",
    "Technology": "Tech.",
    "Technologies": "Techs.",
    "Telecommunication": "Telecomm.",
    "Telecommunications": "Telecomm.",
    "Telephone": "Tel.",
    "Telephones": "Tels.",
    "Telegraph": "Tel.",
    "Telegraphs": "Tels.",
    "Temporary": "Temp.",
    "Temporaries": "Temps.",
    "Township": "Twp.",
    "Townships": "Twps.",
    "Transcontinental": "Transcon.",
    "Transcontinentals": "Transcons.",
    "Transport": "Transp.",
    "Transports": "Transps.",
    "Transportation": "Transp.",
    "Transportations": "Transps.",
    "Trustee": "Tr.",
    "Trustees": "Trs.",
    "Turnpike": "Tpk.",
    "Turnpikes": "Tpks.",
    "Uniform": "Unif.",
    "Uniforms": "Unifs.",
    "University": "Univ.",
    "Universities": "Univs.",
    "Utility": "Util.",
    "Utilities": "Utils.",
    "United States": "U.S.",
    "United States of America": "U.S.",
    "Village": "Vill.",
    "Villages": "Vills.",
    "West": "W.",
    "Western": "W."
}


def extract_first_two_pages(text):
    lines = text.split('\n')
    lines_per_page = 70
    first_two_pages = lines[:lines_per_page*2]

    return '\n'.join(first_two_pages)

def extract_text_from_pdf(pdf_file):
    all_text = ""
    with pdfplumber.open(pdf_file) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Extract header text
            header_text = page.extract_text(x_tolerance=2, y_tolerance=2, layout=True)
            header_text = header_text.split('\n')[0] if header_text else ""
            
            # Extract main page text
            page_text = page.extract_text()
            
            # Check for image-based content by ensuring the presence of text
            if not page_text and page.images:
                page_text = "[Image-based content detected]"
            
            if page_num == 0:
                all_text += f"Header: {header_text.strip()} {page_text.strip() if page_text else ''}\n\n"
            else:
                all_text += f"Header: {header_text.strip()} {page_text.strip() if page_text else ''}\n\n"

    return all_text.strip()


def extract_text_from_docx(docx_file):
    """
    Extract text from a DOCX file, removing duplicate headers.

    Parameters:
        docx_file (str): Path to the DOCX file. 

    Returns:
        str: Extracted text with formatting.
    """
    try:
        full_text = docx2txt.process(docx_file)  # Extract text from the document

        # Replace line breaks with spaces and keep page breaks as new lines
        formatted_text = full_text # Replace line breaks with a space
        formatted_text = formatted_text.replace('\f', '\n')  # Replace page breaks with new lines

        # Split into lines and filter out duplicates
        lines = formatted_text.splitlines()
        unique_lines = list(dict.fromkeys(lines))  # Remove duplicates while preserving order

        # Recombine the unique lines
        cleaned_text = '\n'.join(unique_lines) # Join unique lines with new lines

        return cleaned_text
    except Exception as e:
        return f"Error extracting text: {e}"

# @st.cache_data
# def cached_process_ocr_pdf(file):
#     """Cache the OCR results for each file"""
#     return process_ocr_pdf(file)


def is_image_based_docx(docx_file):
    """
    Check if a DOCX file contains images that might need OCR
    """
    try:
        import docx
        doc = docx.Document(docx_file)
        
        # Check for InlineShapes (images) in the document
        has_images = False
        for paragraph in doc.paragraphs:
            if paragraph.runs:
                for run in paragraph.runs:
                    if run._element.findall('.//pic:pic', doc._element.nsmap):
                        has_images = True
                        break
        
        # Also check if there's very little text content
        text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # If there are images and little text, it might need OCR
        return has_images and len(text_content.strip()) < 200
        
    except Exception as e:
        print(f"Error checking DOCX for images: {e}")
        return False


def convert_docx_to_pdf(docx_file):
    """
    Convert DOCX file to PDF and return as BytesIO object
    """
    try:
        # Create a temporary directory
        with tempfile.TemporaryDirectory() as temp_dir:
            # Create paths for temporary files
            temp_docx = Path(temp_dir) / "temp.docx"
            temp_pdf = Path(temp_dir) / "temp.pdf"
            
            # Save the uploaded file to temporary location
            with open(temp_docx, "wb") as f:
                f.write(docx_file.getvalue())
            
            # Convert to PDF
            convert(str(temp_docx), str(temp_pdf))
            
            # Read the generated PDF into BytesIO
            pdf_bytes = io.BytesIO()
            with open(temp_pdf, "rb") as f:
                pdf_bytes.write(f.read())
            
            # Reset pointer to start
            pdf_bytes.seek(0)
            
            return pdf_bytes
            
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        return None


# def convert_docx_to_pdf(docx_file):
#     """
#     Convert DOCX file to PDF and return as BytesIO object
#     """
#     try:
#         # Create a temporary directory
#         with tempfile.TemporaryDirectory() as temp_dir:
#             # Create paths for temporary files
#             temp_docx = Path(temp_dir) / "temp.docx"
#             temp_pdf = Path(temp_dir) / "temp.pdf"
            
#             # Save the uploaded file to temporary location
#             with open(temp_docx, "wb") as f:
#                 f.write(docx_file.getvalue())
            
#             # Initialize COM for docx2pdf
#             pythoncom.CoInitialize()
            
#             # Convert to PDF
#             convert(str(temp_docx), str(temp_pdf))
            
#             # Read the generated PDF into BytesIO
#             pdf_bytes = io.BytesIO()
#             with open(temp_pdf, "rb") as f:
#                 pdf_bytes.write(f.read())
            
#             # Reset pointer to start
#             pdf_bytes.seek(0)
            
#             return pdf_bytes
            
#     except Exception as e:
#         print(f"Error converting DOCX to PDF: {e}")
#         return None
#     finally:
#         # Uninitialize COM
#         pythoncom.CoUninitialize()

# def extract_text(file):
#     if file is not None:
#         if file.name.endswith('.pdf'):
#             # Create a placeholder for the progress bar
#             progress_placeholder = st.empty()
            
#             # Check if PDF is image-based
#             if is_image_based_pdf(file):
#                 with st.spinner('PDF is image-based. Running OCR... This may take a few minutes...'):
#                     progress_bar = progress_placeholder.progress(0)
                    
#                     # Process OCR extraction
#                     extracted_text = cached_process_ocr_pdf(file)
                    
#                     # Update progress
#                     progress_bar.progress(100)
#                     progress_placeholder.empty()
                    
#                     if extracted_text and any(extracted_text):
#                         return " ".join(extracted_text)
#                     else:
#                         st.warning("OCR extraction failed. The PDF might contain unclear images.")
#                         return None
#             else:
#                 # Normal PDF text extraction
#                 return extract_text_from_pdf(file)
            
#         elif file.name.endswith('.docx'):
#             progress_placeholder = st.empty()
            
#             if is_image_based_docx(file):
#                 with st.spinner('DOCX contains images. Converting to PDF for OCR...'):
#                     progress_bar = progress_placeholder.progress(0)
                    
#                     # Convert DOCX to PDF
#                     pdf_file = convert_docx_to_pdf(file)
#                     extracted_text = cached_process_ocr_pdf(pdf_file)
#                     # Update progress
#                     progress_bar.progress(100)
#                     progress_placeholder.empty()
                    
#                     if extracted_text and any(extracted_text):
#                         return " ".join(extracted_text)
#                     else:
#                         st.warning("DOCX to PDF conversion failed. Please try again.")
#                         return None
#             else:
#                 return extract_text_from_docx(file)
#     return None


def remove_suffix(s):
    if s.endswith("CV"):
        return s[:-3]  # Remove last 2 characters
    elif s.endswith("CR"):
        return s[:-3]  # Remove last 2 characters
    return s
def setOptions (role ):
    if role == "admin":
        return ["Reset Password", "Add User", "Update"]
    elif role == "user":
        return ["Reset Password"]

def text_summarizer_alternate(value):
     # Define the context for the summary
    context = """you are a US lawyer that makes summaries according a specific structure. Here are the instructions :
    the summary can be characterised as a case digest or a case brief. It is a concise restatement of the essential elements of the court''s decision, including:
    1. The procedural context (2 - 4 sentences)
    2. The factual background (2 - 4  sentences)
    3. The legal arguments presented (2 - 4 sentences )
    4. The trial court''s findings  (2 - 4 sentences)
    5. The  court''s decision (2 - 4 sentences)

    Guidelines :
    The summary effectively captures the essence of the decision, highlighting the key legal findings and the rationale for the court''s ruling. It is structured to provide a clear and quick understanding of the outcome and the reasons behind it, which is useful for legal professionals interested into the case.
    The summary needs to be without the titles of the sections , in one block of text. Also you can use roles like : plaintiff, defendant etc... when needed.
    If there is only one plaintiff, defendant or petitioner, then use "defendant" or "plaintiff" or "petitioner" instead of the name.
    
    Also don''t use formulas like : in this case, judgment.
    Don''t use formulas like "In the United States district court for the District of New Jersey", skip it or use "Court."
    Do not need to repeat the name of the case.
    Use "court", instead of "the court", it can be also "family court" etc.., instead of "the family court", basically remove "the" when it is the court or similar.
    Answer in a professional way, don''t invent, stick to the facts.
    if you copy text from the orginal case put into quotes " " .
    Expresion like "filed a motion", can be replaced by "moved to".
    if there are number don''t put them into letters if they are 10 or above, keep them in numbers like 98 or if percentage : 98%.
    if defendant and plaintiff do not start a sentence then they should not be capitalized, even if they are capitlized in the legal decison, don''t capitlize unless it starts a sentence.
    Keep it between 195-375 tokens."""
    
    context = context + """
    In your summary, please ensure the following key aspects are addressed:
    Legal Proceedings Accuracy: Detail the sequence and timing of key legal events, including pre-trial motions, trial proceedings, and post-trial decisions. Highlight any procedural nuances that are critical to understanding the case's context.   
    Detailed Fact Consideration: Include a thorough overview of the factual background of the case. Emphasize key facts that were pivotal in the court's decision-making process, especially those relevant to any summary judgment considerations.

    Clarity and Readability:

    Jurisdictional Clarity: Clearly distinguish between different types of jurisdiction (e.g., subject matter, personal) involved in the case and their implications.
    Legal Standards and Acts: Accurately describe and reference the legal standards and acts applied in the case, ensuring that their relevance and impact on the case are clearly explained.
    Factual vs. Legal Findings: Differentiate between the court's factual findings and legal rulings. Clarify how these findings influenced the case outcome.
    Semantic Precision: Use precise and appropriate legal terminology throughout the summary. Pay close attention to the semantic implications of legal language and ensure that the summary accurately reflects the nuances and complexities of the case.

    Please balance the need for detailed legal analysis with the conciseness expected in a summary, ensuring that the final output is both informative and accessible to legal professionals
    """

    # Call the OpenAI API to generate a summary
    response = client.chat.completions.create(
        model=GPTModel,
        temperature=0.0,
        max_tokens=600,
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": value}
        ]
    )
    
    summary_response = response.choices[0].message.content.strip()
    summary_response = ' '.join(summary_response.splitlines())
    
    summary_response = summary_response.replace("$", "&#36;")
    # print("Request Parameters:", {
    # "model": GPTModel,
    # "temperature": 0.0,
    # "max_tokens": 600,
    # "messages": [
    #     {"role": "system", "content": context},
    #     {"role": "user", "content": value}
    # ]
    # })
    # print("API Response:", response)
    return summary_response 


def abbreviate_title(title):
    words = word_tokenize(title)
    abbreviated_words = []
    for word in words:
        if word in abbrev_dict:
            abbreviated_words.append(abbrev_dict[word])
        elif word.capitalize() in abbrev_dict:
            abbreviated_words.append(abbrev_dict[word.capitalize()])
        else:
            abbreviated_words.append(word)
    return ' '.join(abbreviated_words)

def title(value):

    title_case=""
    
    prompt_title = """
            Give the title of the legal case, take the current title first, then here are the rules : 
            If there are several defendants, just take the first one.
            If it is a person just keep his last name and don't put his first name. 
            If it is a company or organization, it needs to keep the whole name, don't abbreviate anything.
            If it is a State of the USA, just mention the State name.
            
            If the title has the following format plaintiff vs defendant, ONLY if it has that format, then extract the case name from a legal text similar to the following format:
            [Plaintiff Lastname] v. [Defendant Lastname]
            Use Capital letter for the first letter of each word when it makes sense, not needed for capitalization preposition words like "of" or standing letter like v.
            
            If the title has a different format, keep the same format, remove any code or reference, that is not part of the title case and use the previous rules.

            just return the title as an answer nothing else.
    
            """
    title_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.0,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_title},
        {"role": "user", "content": value}
        ]
    )
    print("new2")
    print (title_response.choices[0].message.content)
    
    title_case = title_response.choices[0].message.content
    
    title_case = abbreviate_title(title_case)
    
    print(title_case + ' : NLP')
    return title_case

def Connecticut_summarizer(value):
    summary =""
    first_two_pages = extract_first_two_pages(value)
    name_case = title(first_two_pages)
    
    summary = "CASE: " +  name_case + "  \n"
    prompt_court_option = ("""I will send you a legal decision and you will detect the court that ruled and return it according to a table, just the court name nothing else.
            Here is the structure of table :
            Connecticut Supreme Court
            Connecticut Appellate Court
            Connecticut Superior Court
            District Courts
            Probate Courts
            Family Courts
            Small Claims Courts
            Housing Courts
            Workers' Compensation Commission
            Tax Court
            Juvenile Courts
            Traffic Courts
            Administrative Agencies
            """)
    court_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_court_option},
        {"role": "user", "content": first_two_pages}
        ]
    )
    
    print (court_response.choices[0].message.content)
    court_case = court_response.choices[0].message.content
    
    prompt_num = """
            Give the number of the legal case, 
            just return the number as an answer nothing else
            """
            
    num_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_num},
        {"role": "user", "content": first_two_pages}
        ]
    )
    print (num_response.choices[0].message.content)
    
    num_case = num_response.choices[0].message.content
    
    prompt_judge = "you are a US lawyer, and will read a legal decision and return the name of the judge, only the name, nothing else, in the format : Lastname, Firstname (only first letter of the Firstname). If the case is PER CURIAM, just return : per curiam. If it 's a federal case and district case, replace the first name by : U.S.D.J. Else if it 's a federal case and magistrate case, replace the first name by : U.S.M.J."

    judge_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_judge},
        {"role": "user", "content": first_two_pages}
        ]
    )
            
    judge_name ="" 
            
    if judge_response.choices[0].message.content =="per curiam" :
        judge_name = "per curiam"
    elif "U.S.D.J." in judge_response.choices[0].message.content:
        name = HumanName(judge_response.choices[0].message.content)
        judge_name = name.last + ", U.S.D.J."
        
    elif "U.S.M.J." in judge_response.choices[0].message.content:
        name = HumanName(judge_response.choices[0].message.content)
        judge_name = name.last + ", U.S.M.J."
                
    else:
        name = HumanName(judge_response.choices[0].message.content)
        judge_name = name.last + ", J."  #.capitalize()
            
    
    print (judge_response.choices[0].message.content)
    date_response = client.chat.completions.create(
                model=GPTModel,
                temperature=0.2,
                max_tokens=16,
                messages=[
                    {"role": "system", "content": "When did the judgment happen, if you can't find, look for decided date, also answer with the date only, nothing else, no additional text, just the date, and abreviate the month like this Jan. Feb. March April May June July Aug. Sept. Oct. Nov. Dec."},
                    {"role": "user", "content": value}
                ]
            )

            # Append the court date to the summary
    court_date = date_response.choices[0].message.content.strip()


    prompt_taxonomy = """ I will give you a table with taxonomy , read the legal case, You can use up to three taxonomies. The last topic is usually Civil Appeals or Criminal Appeals. The format is caps and lower case. Just return up to 3 taxonomies, separated by ",", example :  Wrongful Death, Civil Rights, Civil Appeals. Here is the table of taxonomy :
            Administrative Law
            Admiralty
            Antitrust
            Banking and Finance Laws
            Bankruptcy
            Civil Procedure
            Civil Rights
            Commercial Law
            Constitutional Law
            Consumer Protection
            Contracts
            Contractual Disputes
            Corporate Entities
            Corporate Governance
            Creditors' and Debtors' Rights
            Criminal Law
            Damages
            Personal Injury
            Dispute Resolution
            Education Law
            Elder Law
            Employment Benefits
            Employment Litigation
            Employment Compliance
            Employment Litigation
            Entertainment and Sports Law
            Environmental Law
            Evidence
            Family Law
            Government
            Health Care Law
            Immigration Law
            Insurance Law
            Intellectual Property
            Internet Law
            Judges
            Legal Ethics and Attorney Discipline
            Legal Malpractice
            Labor Law
            Employment Benefits
            Employment Compliance
            Employment Litigation
            Land Use and Planning
            Landlord/Tenant
            Mass Tort Claims
            Motor Vehicle Torts
            Toxic Torts
            Business Torts
            Damages
            Medical Malpractice
            Motor Vehicle Torts
            Products Liability
            Public Records
            Public Utilities
            Real Estate
            Securities
            Tax
            Telecommunications
            Transportation
            Trusts and Estates
            Wrongful Death
            Civil Appeals
            Criminal Appeals
            """
    taxonomy_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_taxonomy},
        {"role": "user", "content": value}
        ]
    )
    print (taxonomy_response.choices[0].message.content)
    practice_area = taxonomy_response.choices[0].message.content
    
    
    prompt_practice = """
            I will give you a table with industries, read the legal case, give the industry of the legal case, just return the industry as an answer nothing else. Here is the table of industries:
            Accounting		
            Advertising		
            Aerospace		
            Agriculture 		
            Autokotive		
            Biotechnology		
            Brokerage		
            Call Centers		
            Cargo and Shipping		
            Chemicals and Materials		
            Construction		
            Consumer Products		
            Defense		
            Ditribution and Wholesale		
            E-Commerce		
            Education		
            Electronics		
            Energy		
            Entertainment and Leisure		
            Executive Search		
            Federal Government		
            Food and Beverage		
            Hardware (Computer)		
            Health Care		
            Hospitatlity and Lodging		
            Insurance		
            Investments and Investment Advisory		
            Legal Services		
            Manufacturing		
            Mining and Resources		
            Non-Profit		
            Pharmaceuticals		
            Real Estate		
            Revruitment and Staffing		
            Retail		
            Software		
            State and Local Government		
            Technology Media Telecom		
            Transportation		
            Travel and Tourism		
            """
                 
    practice_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_practice},
        {"role": "user", "content": value}
        ]
    )
    print (practice_response.choices[0].message.content)
    
    practice_case = practice_response.choices[0].message.content
    prompt_title = "you are a US lawyer, and will read a legal decision and return the title of the case, only the title, nothing else, the title should describe in a sentence the case without mentioning the plaintiff and the defendants."

    title_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_title},
        {"role": "user", "content": first_two_pages}
        ]
    )
    print (title_response.choices[0].message.content)
    title_case = title_response.choices[0].message.content
    summary = f"""
CASE: {name_case} \n
COURT: {court_case} \n
DOC NO: {num_case} \n
COURT OPINION BY: {judge_name} \n
DATE: {court_date} \n
PAGES: {page_count} \n

{text_summarizer_alternate(value)}

TITLE: {title_case}               \n
LITIGATION: Yes  \n
PRACTICE AREA: {practice_area} \n
INDUSTRY: {practice_case} \n
ROLE:  


"""
    return summary

def Texas_summarizer(value):
    summary =""
    
    # Display the generated summary
    summary = text_summarizer_alternate(value)
    first_two_pages = extract_first_two_pages(value)
    print(summary)
    
    summary = summary.replace("District Court", "district court")
    
    prompt_taxonomy = """ I will give you a table with taxonomy , read the legal case, You can use up to three taxonomies. The last topic is usually Civil Appeals or Criminal Appeals. The format is caps and lower case. Just return up to 3 taxonomies, separated by "│", example :  Wrongful Death│Civil Rights│Civil Appeals. Here is the table of taxonomy :
            Administrative Law
            Admiralty
            Antitrust
            Banking and Finance Laws
            Bankruptcy
            Civil Procedure
            Civil Rights
            Commercial Law
            Constitutional Law
            Consumer Protection
            Contracts
            Contractual Disputes
            Corporate Entities
            Corporate Governance
            Creditors' and Debtors' Rights
            Criminal Law
            Damages
            Personal Injury
            Dispute Resolution
            Education Law
            Elder Law
            Employment Benefits
            Employment Litigation
            Employment Compliance
            Employment Litigation
            Entertainment and Sports Law
            Environmental Law
            Evidence
            Family Law
            Government
            Health Care Law
            Immigration Law
            Insurance Law
            Intellectual Property
            Internet Law
            Judges
            Legal Ethics and Attorney Discipline
            Legal Malpractice
            Labor Law
            Employment Benefits
            Employment Compliance
            Employment Litigation
            Land Use and Planning
            Landlord/Tenant
            Mass Tort Claims
            Motor Vehicle Torts
            Toxic Torts
            Business Torts
            Damages
            Medical Malpractice
            Motor Vehicle Torts
            Products Liability
            Public Records
            Public Utilities
            Real Estate
            Securities
            Tax
            Telecommunications
            Transportation
            Trusts and Estates
            Wrongful Death
            Civil Appeals
            Criminal Appeals
            """
    taxonomy_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_taxonomy},
        {"role": "user", "content": value}
        ]
    )
    print (taxonomy_response.choices[0].message.content)
    
    summary = taxonomy_response.choices[0].message.content + "  \n" + summary
    summary = summary + "  \n" + title(first_two_pages)
    
    
    prompt_court_option = ("""I will send you a legal decision and you will detect the court that ruled and return it according to a table, just the court name nothing else.
            Here is the structure of table :
            Fifth Circuit
            Supreme Court of Texas
            Court of Criminal Appeals
            First Court of Appeals
            Second Court of Appeals 
            ...
            Fourteenth Court of Appeals
            [District number] Court of Appeals                
            """)
    court_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": prompt_court_option},
        {"role": "user", "content": first_two_pages}
        ]
    )
    print (court_response.choices[0].message.content)
    summary = summary + ", " + court_response.choices[0].message.content
    
    case_number = ('I will send you a legal decision and you will detect the case number and return it, just the case number nothing else ')
            
    case_number_response = client.chat.completions.create(
    model = GPTModel,
    temperature = 0.2,
    max_tokens = 600,
    messages = [
        {"role": "system", "content": case_number},
        {"role": "user", "content": first_two_pages}
        ]
    )
    print (case_number_response.choices[0].message.content)
    
    case_num = remove_suffix(case_number_response.choices[0].message.content)
    summary = summary + ", " + case_num
    
    # Extract the court date
    date_response = client.chat.completions.create(
        model=GPTModel,
        temperature=0.2,
        max_tokens=16,
        messages=[
            {"role": "system", "content": "When did the judgment happen, if you can't find, look for decided date, also answer with the date only, nothing else, no additional text, just the date, with this format 01/17/2021"},
            {"role": "user", "content": value}
        ]
    )

    # Append the court date to the summary
    court_date = date_response.choices[0].message.content.strip()
    summary = summary + ", " + court_date
               
        
    return summary
    

# Function to validate if the input is a positive integer
def is_positive_integer(value):
    try:
        value = int(value)
        return value > 0
    except ValueError:
        return False

def initialize_session_state():
    if 'extracted_text' not in st.session_state:
        st.session_state.extracted_text = None
    if 'processing_complete' not in st.session_state:
        st.session_state.processing_complete = False
    
# Define the Streamlit app
def main():
    check_openai_key(OPENAI_API_KEY)
    ensure_nltk_data()
    global page_count
    st.image('MESJ.jpg')

    with open('config.YAML') as file:
        config = yaml.load(file, Loader=SafeLoader)
    with open('cfg1.YAML') as file:
        roles_config = yaml.load(file, Loader=SafeLoader)

    authenticator = stauth.Authenticate(
        config['credentials'],
        config['cookie']['name'],
        config['cookie']['key'],
        config['cookie']['expiry_days']
    )

    # Initialize authentication status if not already in session state
    if 'authentication_status' not in st.session_state:
        st.session_state.authentication_status = None

    # Login Widget
    authenticator.login('main')

    # Authentication Status Handling
    if st.session_state["authentication_status"]:
        # Successful Authentication
        name = st.session_state["name"]
        username = st.session_state["username"]
        
        try:
            # Determine user role from configuration
            role = roles_config["usernames"][username]["role"]
        except KeyError:
            # Fallback to default user role if not found
            role = "user"

        # Authenticated User Interface
        authenticator.logout("Logout", "sidebar")
        st.sidebar.title(f"Welcome {name}")
        
        # Application Mode Selection
        app_mode = st.sidebar.selectbox("Choose your preference:", ["Legal Decision Summarizer", "Newsletter Quotes"])

        # Options Selection Based on User Role
        selected_option = st.sidebar.selectbox("Options", setOptions(role))

        # Expanded Options Handling
        with st.expander(f"{selected_option}"):
            # Perform actions based on the selected option
            if selected_option == "Reset Password":
                try:
                    if authenticator.reset_password(username):
                        st.success('Password modified successfully')
                        # Update configuration file
                        with open('config.YAML', 'w') as file:
                            yaml.dump(config, file, default_flow_style=False)
                except Exception as e:
                    st.error(f"Password reset error: {e}")
            elif selected_option == "Add User":
                try:
                    email_of_registered_user, username_of_registered_user, name_of_registered_user = authenticator.register_user(
                        pre_authorized=config['pre-authorized']['emails'],
                        fields={
                            'Form name': 'Register user',
                            'Email': 'Email',
                            'Username': 'Username',
                            'Password': 'Password',
                            'Repeat password': 'Repeat password',
                            'Register': 'Register'
                        }
                    )
                    
                    if username_of_registered_user:
                        st.session_state["username_of_registered_user"] = username_of_registered_user
                        st.success('User registered successfully')
                        # Separate role and state assignment logic
                        role = st.selectbox("Select a role", options=["admin", "user"])
                        
                        if role == 'user':
                            chosen_states = st.multiselect("Select states", options=["New Jersey", "Texas", "Connecticut"])
                            
                            if chosen_states:
                                with open('cfg1.YAML', 'r+') as file:
                                    role_data = yaml.safe_load(file)
                                    
                                    role_data["usernames"][username_of_registered_user] = {
                                        "role": role,
                                        "states": chosen_states
                                    }
                                    
                                    file.seek(0)
                                    yaml.dump(role_data, file)
                                    file.truncate()
                        
                        # Update config file
                        with open('config.YAML', 'w') as file:
                            yaml.dump(config, file, default_flow_style=False)
                        
                        st.success('User registered successfully')
                    if "username_of_registered_user" in st.session_state.keys() and st.session_state["username_of_registered_user"]:
                        role='admin'
                        role = st.selectbox("select a role", options=["admin", "user"])
                        if role =='user':
                            chosen_states = st.multiselect ("select the states" , options = ["New Jersey", "Texas", "Connecticut"])
                            if role and chosen_states:
                                with open('cfg1.YAML', 'r') as file:
                                    role_data = yaml.safe_load(file)
                                    if st.session_state["username_of_registered_user"] not in role_data['usernames'].keys() or role_data["usernames"][st.session_state["username_of_registered_user"]]==None :
                                        role_data["usernames"][st.session_state["username_of_registered_user"]]= dict()
                                    role_data["usernames"][st.session_state["username_of_registered_user"]]["role"] = role
                                    role_data["usernames"][st.session_state["username_of_registered_user"]]["states"] = chosen_states
                                with open('cfg1.YAML', 'w') as file:
                                    yaml.dump(role_data, file)
                                    file.close()
                        
                except Exception as e:
                    st.error(e)
                with open('config.YAML', 'w') as file:
                        yaml.dump(config, file, default_flow_style=False)
                
            elif selected_option == "Update":
                if st.session_state["authentication_status"]:   
                    try:
                        with open('config.YAML', 'r') as file:
                                usernames_data = yaml.safe_load(file)["credentials"]["usernames"]
                                usernames_= usernames_data.keys()
                        username_updated= st.selectbox('username',usernames_)
                        if username_updated : 
                            st.session_state["username_updated"] = username_updated
                        if authenticator.update_user_details(username_updated):
                            st.success('Entries updated successfully')
                        with open('cfg1.YAML', 'r') as file:
                            role_data = yaml.safe_load(file)
                        if "username_updated" in st.session_state.keys() and st.session_state["username_updated"] and role_data["usernames"][st.session_state["username_updated"]]['role']== 'user':                            
                                chosen_states = st.multiselect ("select the states" , options = ["New Jersey", "Texas", "Connecticut"])
                                if chosen_states:
                                    
                                    if st.session_state["username_updated"] not in role_data['usernames'].keys() or role_data["usernames"][st.session_state["username_updated"]]==None :
                                        role_data["usernames"][st.session_state["username_updated"]]= dict()
                                    role_data["usernames"][st.session_state["username_updated"]]["states"] = chosen_states
                                    with open('cfg1.YAML', 'w') as file:
                                        yaml.dump(role_data, file)
                                        file.close()
                    except Exception as e:
                        st.error(e)
                    with open('config.YAML', 'w') as file:
                        yaml.dump(config, file, default_flow_style=False)
                            
        if app_mode == "Legal Decision Summarizer":
            st.title("Legal Decision Summarizer")
            
            initialize_session_state()
            
            choice1 = st.radio("How would you like to provide the legal decision?", ('Copy-Paste Text', 'Upload Document'))
            
            # Initialize variables
            show_additional_inputs = True
            user_input = None
            first_two_pages = None

            if choice1 == 'Copy-Paste Text':
                user_input = st.text_area("Enter legal decision:", height=150)
                if user_input:
                    first_two_pages = extract_first_two_pages(user_input)
                    st.session_state.extracted_text = user_input
                    st.session_state.processing_complete = True

            elif choice1 == 'Upload Document':
                user_file_input = st.file_uploader("Upload your document", type=["pdf", "docx"])

                if user_file_input is not None:
                    # Create progress placeholder
                    progress_placeholder = st.empty()
                    status_placeholder = st.empty()

                    if not st.session_state.processing_complete:
                        if user_file_input.name.endswith('.pdf'):
                            status_placeholder.info("Processing PDF... Please wait...")
                            progress_bar = progress_placeholder.progress(0)

                            # Process PDF
                            combined_text = None
                            try:
                                if is_image_based_pdf(user_file_input):
                                    status_placeholder.warning("PDF is image-based. Running OCR... This may take a few minutes...")
                                    progress_bar.progress(25)
                                    
                                    # Process OCR
                                    extracted_text = process_ocr_pdf(user_file_input)
                                    progress_bar.progress(75)
                                    
                                    if extracted_text and any(extracted_text):
                                        combined_text = " ".join(extracted_text)
                                    else:
                                        st.error("OCR extraction failed. The PDF might contain unclear images.")
                                        show_additional_inputs = False
                                else:
                                    status_placeholder.info("Extracting text from PDF...")
                                    progress_bar.progress(50)
                                    combined_text = extract_text_from_pdf(user_file_input)
                                    progress_bar.progress(90)

                            except Exception as e:
                                st.error(f"Error processing PDF: {str(e)}")
                                show_additional_inputs = False

                        elif user_file_input.name.endswith('.docx'):
                            status_placeholder.info("Processing DOCX... Please wait...")
                            progress_bar = progress_placeholder.progress(0)
                            combined_text = None
                            try:
                                if is_image_based_docx(user_file_input):
                                    status_placeholder.warning("DOCX is image-based. Running OCR... This may take a few minutes...")
                                    progress_bar.progress(25)
                                    pdf_file = convert_docx_to_pdf(user_file_input)
                                    extracted_text = process_ocr_pdf(pdf_file)
                                    progress_bar.progress(75)
                                    
                                    if extracted_text and any(extracted_text):
                                        combined_text = " ".join(extracted_text)
                                    else:
                                        st.error("OCR extraction failed. The DOCX might contain unclear images.")
                                        show_additional_inputs = False
                                            
                                # combined_text = extract_text_from_docx(user_file_input)
                                # progress_bar.progress(90)
                                else:
                                    status_placeholder.info("Extracting text from DOCX...")
                                    progress_bar.progress(50)
                                    combined_text = extract_text_from_docx(user_file_input)
                                    progress_bar.progress(90)
                            except Exception as e:
                                st.error(f"Error processing DOCX: {str(e)}")
                                show_additional_inputs = False

                        # Process the extracted text
                        if combined_text:
                            if len(combined_text.strip()) < 200:
                                st.error(f"Uploaded {user_file_input.name.split('.')[-1].upper()} contains very little text. Please upload a valid document.")
                                show_additional_inputs = False
                            else:
                                first_two_pages = extract_first_two_pages(combined_text)
                                user_input = combined_text
                                st.session_state.extracted_text = combined_text
                                st.session_state.processing_complete = True
                        else:
                            st.error(f"Could not extract text from the {user_file_input.name.split('.')[-1].upper()} file.")
                            show_additional_inputs = False

                        # Clean up progress indicators
                        progress_placeholder.empty()
                        status_placeholder.empty()

                    else:
                        # Use cached results
                        user_input = st.session_state.extracted_text
                        first_two_pages = extract_first_two_pages(user_input)

                else:
                    st.warning("No file uploaded. Please upload a document.")
                    show_additional_inputs = False

            # Only show additional inputs if we have valid text and processing is complete
            if show_additional_inputs and st.session_state.processing_complete:
                if role == "user":
                    try:
                        states = roles_config["usernames"][username]["states"]
                    except:
                        states = []
                else:
                    states = ["New Jersey", "Texas", "Connecticut"]
                
                state = st.selectbox("Select a US State:", states)
                
                if state != "Texas":
                    page_count_input = st.text_input("Page count:", value="1")
                    if is_positive_integer(page_count_input):
                        page_count = int(page_count_input)
                    else:
                        st.error("Please enter a valid positive integer for the page count.")
                        page_count = None
                else:
                    page_count = None

                if st.button("Summarize"):
                    if state == "New Jersey":

                        # Display the generated summary
                        summary = text_summarizer_alternate(user_input) 
                    
                        
                        print(summary)
                        
                        summary = summary.replace("District Court", "district court")
                        st.subheader("Summary:")

                        # Type of case federal or State
                        federal_response = client.chat.completions.create(
                            model=GPTModel,
                            temperature=0.2,
                            max_tokens=16,
                            messages=[
                                {"role": "system", "content": """
                                
                                Determine if the legal case, if related to a state or federal case, the federal cases are these 
                                Bankr. D.N.J. (U.S Bankruptcy Court) 6
                                D.N.J. (U.S. District Court) - 7
                                3d Cir. (Third Circuit) – 8
                                
                                If that's a federal case just return Federal, nothing else, if it's a state just retrun State nothing else.
                                """},
                                {"role": "user", "content": user_input}
                            ]
                        )

                        # Append the court date to the summary
                        court_type = federal_response.choices[0].message.content.strip()


                        # Extract the court date
                        date_response = client.chat.completions.create(
                            model=GPTModel,
                            temperature=0.2,
                            max_tokens=16,
                            messages=[
                                {"role": "system", "content": "Check filed date, usually it is at the top of the document, American date format, also answer with the date only, nothing else, no additional text, just the date, and abreviate the month like this Jan. Feb. March April May June July Aug. Sept. Oct. Nov. Dec."},
                                {"role": "user", "content": user_input}
                            ]
                        )

                        # Append the court date to the summary
                        court_date = date_response.choices[0].message.content.strip()
                        
                        if court_type =="Federal":
                            summary = summary + " [Filed " + court_date + "]"
                        
                        # judge
                        prompt_judge = "you are a US lawyer, and will read a legal decision and return the name of the judge, only the name, nothing else, in the format : Lastname, Firstname (only first letter of the Firstname). If the case is PER CURIAM, just return : per curiam. If it 's a federal case and district case, replace the first name by : U.S.D.J. Else if it 's a federal case and magistrate case, replace the first name by : U.S.M.J."

                        judge_response = client.chat.completions.create(
                        model = GPTModel,
                        temperature = 0.0,
                        max_tokens = 600,
                        messages = [
                            {"role": "system", "content": prompt_judge},
                            {"role": "user", "content": user_input}
                            ]
                        )
                        
                        judge_name ="" 
                        
                        if judge_response.choices[0].message.content =="per curiam" :
                            judge_name = "per curiam"
                        elif "U.S.D.J." in judge_response.choices[0].message.content:
                            name = HumanName(judge_response.choices[0].message.content)
                            judge_name = name.last + ", U.S.D.J."
                            
                        elif "U.S.M.J." in judge_response.choices[0].message.content:
                            name = HumanName(judge_response.choices[0].message.content)
                            judge_name = name.last + ", U.S.M.J."
                            
                        else:
                            name = HumanName(judge_response.choices[0].message.content)
                            judge_name = name.last + ", J."  #.capitalize()
                        
                        summary = " (" + judge_name + ") (" + str(page_count) + " pp.) "  + summary 
                        print (judge_response.choices[0].message.content)
                        
                        # court option
                        
                        courts = {
                                'N.J.': 1,
                                'N.J. Super. App. Div.': 2,
                                'N.J. Super. Law Div.': 3,
                                'N.J. Super. Ch. Div.': 4,
                                'Tax Ct.': 5,
                                'Bankr. D.N.J.': 6,
                                'D.N.J.': 7,
                                '3d Cir.': 8
                                }
                        courts_inverted = {value: key for key, value in courts.items()}

                        
                        prompt_court_option = ('I will send you a legal decision and you have to select one of these court option, just return the corresponding number, nothing else, here are the court option :' 
                            'N.J. Sup. Ct. (Supreme Court) - 1 '
                            'N.J. Super. App. Div. (Appellate Division) 2 '
                            'N.J. Super. Law Div. – (Law Division) (Civil and Criminal) 3 '
                            'N.J. Super. Ch. Div. (Chancery Division) (General Equity and Family) -4 '
                            'Tax Ct. – (Tax Court) - 5 '
                            'Bankr. D.N.J. (U.S Bankruptcy Court 6 '
                            'D.N.J. (U.S. District Court) - 7 '
                            '3d Cir. (Third Circuit) - 8 ')
                        
                        court_response = client.chat.completions.create(
                        model = GPTModel,
                        temperature = 0.2,
                        max_tokens = 600,
                        messages = [
                            {"role": "system", "content": prompt_court_option},
                            {"role": "user", "content": first_two_pages}
                            ]
                        )
                        print (court_response.choices[0].message.content)
                        summary = courts_inverted[int(court_response.choices[0].message.content)] + " "  + summary
                        
                        title_case = (f"*{title(first_two_pages)}*")
                        
                        
                        summary = title_case + ", "  + summary 
                        
                        # taxonomy
                        prompt_taxonomy = """ I will give you a table with taxonomy , read the legal case, just return the corresponding number , nothing else. here is the table :
                            NJ topic #	NJ Taxonomy Topics
                            01	Administrative Law
                            54	Admiralty
                            59	Antitrust
                            06	Banking and Finance Laws
                            42	Bankruptcy
                            07	Civil Procedure
                            46	Civil Rights
                            08	Commercial Law
                            10	Constitutional Law
                            09	Consumer Protection
                            11	Contracts; Contractual Disputes
                            12	Corporate Entities; Corporate Governance
                            15	Creditors' and Debtors' Rights
                            14	Criminal Law
                            31	Damages; Personal Injury
                            03	Dispute Resolution
                            16	Education Law
                            60	Elder Law
                            39	Employment Benefits; Employment Litigation
                            55	Entertainment and Sports Law
                            17	Environmental Law
                            19	Evidence
                            20	Family Law
                            21	Government
                            22	Health Care Law
                            51	Immigration Law
                            23	Insurance Law
                            53	Intellectual Property
                            61	Internet Law
                            48	Judges
                            04	Judges; Legal Ethics and Attorney Discipline; Legal Malpractice
                            56	Labor Law; Employment Benefits
                            25	Labor Law; Employment Compliance; Employment Litigation
                            26	Land Use and Planning
                            27	Landlord/Tenant
                            36	Mass Tort Claims; Motor Vehicle Torts; Toxic Torts; Business Torts; Damages
                            29	Medical Malpractice
                            05	Motor Vehicle Torts
                            32	Products Liability
                            52	Public Records
                            37	Public Utilities
                            34	Real Estate
                            50	Securities
                            35	Tax
                            57	Telecommunications
                            49	Transportation
                            38	Trusts and Estates
                            40	Wrongful Death
                            """

                        taxonomy_response = client.chat.completions.create(
                        model = GPTModel,
                        temperature = 0.2,
                        max_tokens = 600,
                        messages = [
                            {"role": "system", "content": prompt_taxonomy},
                            {"role": "user", "content": user_input}
                            ]
                        )
                        print (taxonomy_response.choices[0].message.content)
                        summary = taxonomy_response.choices[0].message.content + "-" + court_response.choices[0].message.content + "-XXXX " + summary
                        
                        hash_table = {
                            "01": "Administrative Law",
                            "54": "Admiralty",
                            "59": "Antitrust",
                            "06": "Banking and Finance Laws",
                            "42": "Bankruptcy",
                            "07": "Civil Procedure",
                            "46": "Civil Rights",
                            "08": "Commercial Law",
                            "10": "Constitutional Law",
                            "09": "Consumer Protection",
                            "11": "Contracts; Contractual Disputes",
                            "12": "Corporate Entities; Corporate Governance",
                            "15": "Creditors' and Debtors' Rights",
                            "14": "Criminal Law",
                            "31": "Damages; Personal Injury",
                            "03": "Dispute Resolution",
                            "16": "Education Law",
                            "60": "Elder Law",
                            "39": "Employment Benefits; Employment Litigation",
                            "55": "Entertainment and Sports Law",
                            "17": "Environmental Law",
                            "19": "Evidence",
                            "20": "Family Law",
                            "21": "Government",
                            "22": "Health Care Law",
                            "51": "Immigration Law",
                            "23": "Insurance Law",
                            "53": "Intellectual Property",
                            "61": "Internet Law",
                            "48": "Judges",
                            "04": "Judges; Legal Ethics and Attorney Discipline; Legal Malpractice",
                            "56": "Labor Law; Employment Benefits",
                            "25": "Labor Law; Employment Compliance; Employment Litigation",
                            "26": "Land Use and Planning",
                            "27": "Landlord/Tenant",
                            "36": "Mass Tort Claims; Motor Vehicle Torts; Toxic Torts; Business Torts; Damages",
                            "29": "Medical Malpractice",
                            "05": "Motor Vehicle Torts",
                            "32": "Products Liability",
                            "52": "Public Records",
                            "37": "Public Utilities",
                            "34": "Real Estate",
                            "50": "Securities",
                            "35": "Tax",
                            "57": "Telecommunications",
                            "49": "Transportation",
                            "38": "Trusts and Estates",
                            "40": "Wrongful Death"
                        }
                        
                        legal_category = hash_table.get(taxonomy_response.choices[0].message.content, "Unknown code").upper()
                        
                        st.markdown(f"**{legal_category}**")
                        st.write(summary)
                    elif state =="Connecticut":
                        st.subheader("Summary:")
                        st.write(Connecticut_summarizer(user_input))
                    elif state == "Texas":
                        st.subheader("Summary:")
                        st.write(Texas_summarizer(user_input))
                    else:
                        st.warning("Please select a state before clicking 'Summarize'.")
                
        elif app_mode == "Newsletter Quotes":

            def process_single_link(item):
                """
                Process a single link with error handling and concurrent execution
                
                Args:
                    item (dict): Dictionary containing link and other item details
                
                Returns:
                    dict or None: Processed newsletter data or None if processing fails
                """
                try:
                    # Attempt to scrape web content
                    web_content = scrap_web(item['link'])
                    if web_content is None:
                        st.warning(f"Failed to scrape content from {item['link']}")
                        return None

                    # Extract newsletter information
                    newsletter_topic = get_topic_newsletter(web_content)
                    newsletter_data = newsletter(web_content)
                    newsletter_background = get_newsletter_background(web_content)

                    if newsletter_data is None:
                        return None

                    # Process people quotes
                    people_quotes = newsletter_data['newsletter']['people']
                    background = newsletter_background.get('background', 'No background available')
                    quoted = newsletter_data.get('quoted', 'No quotes available')
                    extracted_topic = newsletter_topic.get('topic', 'No topic found from web content')
                    
                    extracted_people_quotes = [
                        {
                            'name': person['name'],
                            'quote': '\n'.join([f'"{quote}"' for quote in person["quote"]])
                        } for person in people_quotes
                    ]

                    # Format date
                    formatted_date = format_date_and_info(item['date'])

                    # Construct final data dictionary
                    return {
                        'topic': extracted_topic,
                        'background': background,
                        'people_quotes': extracted_people_quotes,
                        'quoted': quoted,
                        'link': item['link'],
                        'date': formatted_date,
                        'branch_head': item['branch_head']
                    }
                
                except KeyError as e:
                    st.warning(f"Error processing data for {item['link']}: Missing key {e}. Skipping...")
                except Exception as e:
                    st.warning(f"An error occurred with {item['link']}: {e}. Skipping...")
                
                return None

            def process_data(uploaded_file):
                """
                Process newsletter quotes data from an uploaded file concurrently
                
                Args:
                    uploaded_file (file): Excel file to be processed
                
                Returns:
                    list: Processed newsletter items
                """
                # Read the Excel file
                df = pd.read_excel(uploaded_file, header=None)
                df.columns = ['A', 'B', 'C', 'D', 'E', 'F']
                
                # Filter and prepare items for processing
                results = list(filter(None, df.apply(process_row, axis=1)))
                
                # Process links concurrently
                with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
                    # Submit all tasks and collect futures
                    futures = [executor.submit(process_single_link, item) for item in results]
                    
                    # Collect results as they complete
                    all_items = [
                        future.result() 
                        for future in concurrent.futures.as_completed(futures) 
                        if future.result() is not None
                    ]
                
                return all_items


            # Title of the app
            st.title('Newsletter Quotes')

            # Uploading the file
            # uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")

            # Check if we already have the processed data in session state
            if 'processed_data' not in st.session_state:
                st.session_state['processed_data'] = None
                
            if 'file_uploader_key' not in st.session_state:
                st.session_state.file_uploader_key = 0  
                    
            if 'downloaded' not in st.session_state:
                st.session_state['downloaded'] = False
                
            if st.session_state['processed_data'] is None:
                uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx", key=st.session_state.file_uploader_key)

                # If a file is uploaded and not already processed
                if uploaded_file is not None:
                    with st.spinner("Processing..."):
                        st.session_state['processed_data'] = process_data(uploaded_file)

            # If the data is processed
            if st.session_state['processed_data']:
                docx_path = create_docx(st.session_state['processed_data'])
                with open(docx_path, "rb") as file:
                    docx_data = file.read()

                # Only show the download button if the file hasn't been downloaded yet
                if not st.session_state['downloaded']:
                    # Display the download button
                    if st.download_button(
                        label="Download DOCX File",
                        data=docx_data,
                        file_name="newsletter_output.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        # Once the file is downloaded, set 'downloaded' to True
                        st.session_state['downloaded'] = True
                        
                process_button_placeholder = st.empty()
                
                if st.session_state['downloaded']:
                    if process_button_placeholder.button("Process New File"):
                        st.session_state['processed_data'] = None
                        st.session_state['downloaded'] = False  # Reset the download state
                        st.session_state.file_uploader_key += 1  # Increment the file uploader key to reset the uploader
                        process_button_placeholder.empty()  # This removes the button after click
                        st.rerun() 

if __name__ == "__main__":
    main()

