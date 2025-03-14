import pandas as pd
import os
import json
import re
from docx import Document
from openai import OpenAI
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from datetime import datetime
import streamlit as st
import fitz  # PyMuPDF
import easyocr
from PIL import Image
import numpy as np
from concurrent.futures import ThreadPoolExecutor
from dotenv import load_dotenv

load_dotenv()
OPENAI_API_KEY= os.getenv("OPENAI_API_KEY")

client = OpenAI(
    api_key = OPENAI_API_KEY,
)

# page_count= None

GPTModelLight = "gpt-4o-mini"
GPTModel = "gpt-4o"



def format_date_and_info(date_str):
    try:
        date_part, info_part = date_str.split("; ", 1)
        # Parse the date with the format MM/DD/YY
        date_obj = datetime.strptime(date_part, "%m/%d/%y")
        # Format the date as 'Month Day, Year'
        formatted_date = date_obj.strftime("%B %d, %Y")  # For example: 'September 07, 2024'
        
        # Extract the location part within parentheses
        if "(" in info_part and ")" in info_part:
            info_main, location = info_part.split("(", 1)
            location = location.rstrip(")")
            formatted_info = f"{info_main.strip()} ({location.strip()})"
        else:
            formatted_info = info_part.strip()
        
        return f"{formatted_date}; {formatted_info}"
    except ValueError as e:
        return date_str
    
    except Exception as e:
        # Handle any error that occurs during date parsing or formatting
        # print(f"Error formatting date: {e}")
        return date_str

        
def process_row(row):
    if pd.isna(row['B']):  # Skip empty rows
        return None
    
    name_role = row['B'].split(', ', 1)
    name = name_role[0]
    role = name_role[1] if len(name_role) > 1 else ''
    branch_head = row['C']
    date_source = row['D'].split(', ', 1)
    date = date_source[0]
    
    return {
        'name': name,
        'role': role,
        'branch_head':branch_head,
        'date': date,
        'link': row['E'],
        'info': row['F']
    }

def newsletter(data):
    # Define the context for the summary
    context = f""" you are a newsletter analyzer. you will be provided with extracted contents of a webpage. your role is to thoroughly analyze the contents of that webpage
    and only focus on the part where some people has quoted their view or their words along with their name that lies in the entire web content.
    After getting those names also get the role of those personnel and collectively store in a "quoted" section of json.
    Here is the web content extracted from the webpage: {data}. 
    
    Return the results in a proper JSON format directly like this:
    
    {{
        "newsletter": \u0028
            "people": [
            \u0028
                "name": "Person 1",
                "quote": [
                "Quote by person 1",
                "another quote by same person (if any)"
                ]
            \u0029,
            \u0028
                "name": "Person 2",
                "quote": [
                "Quote by person 2",
                "another quote by same person (if any)"
                ]
            \u0029
            // more if any...
            ]
        \u0029,
        "quoted": "Person 1 - Role 1, Person 2 - Role 2, Person 3 - Role 3,......"
    }}

    """

    # Call the OpenAI API to generate a summary
    response = client.chat.completions.create(
        model=GPTModel,
        temperature=0.3,
        max_tokens=2000,
        messages=[
            {"role": "system", "content": context},
            # {"role": "user", "content": data}
        ]
    )
    
    response_content = response.choices[0].message.content
    print({"response_content":response_content})
    cleaned_response = response_content.replace('```json\n', "").replace('\n```', "")
    print({"cleaned":cleaned_response})
    # Parse the JSON response
    try:
        structured_data = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        print("Failed to decode JSON:", e)
        return None
    print(structured_data)
    return structured_data


def get_topic_newsletter(data):
    context = f"""
    I have an article related to drugs. Please generate a topic following this structure: [Main Subject]: [Key Elements].
    Here is an example :
    **Topic:** "Xylazine: Danger Warning: Pennsylvania"
    **Article Summary:** "Pennsylvania's Department of Health Secretary Bogen unveils a new wound care initiative aimed at addressing injuries associated with the use of xylazine, an emerging and dangerous drug."
    
    Here is the article summary: [{data}]
    
    Directly return topic in str in json format:
    {{
    "topic": "topic related to article."
    }}
    """
    # Call the OpenAI API to generate a summary
    response = client.chat.completions.create(
        model=GPTModel,
        temperature=0.0,
        max_tokens=600,
        messages=[
            {"role": "system", "content": context},
            # {"role": "user", "content": data}
        ]
    )
    
    response_content = response.choices[0].message.content
    cleaned_response = response_content.replace('```json\n', "").replace('\n```', "")
    

    # Parse the JSON response
    try:
        structured_data = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        print("Failed to decode JSON:", e)
        return None
    print(structured_data)
    return structured_data


def get_newsletter_background(data):
    context = f"""  you will be provided with extracted contents of a webpage. your role is to return first two paragraph as background with very little 
    or no modification at all.
    here is the web content:({data})
    Return background in json format directly:
    {{
    "background": "First two paragraphs of the web content for background."
    }}

    """
    # Call the OpenAI API to generate a summary
    response = client.chat.completions.create(
        model=GPTModel,
        temperature=0.0,
        max_tokens=1000,
        messages=[
            {"role": "system", "content": context},
            # {"role": "user", "content": data}
        ]
    )
    
    response_content = response.choices[0].message.content
    cleaned_response = response_content.replace('```json\n', "").replace('\n```', "")
    

    # Parse the JSON response
    try:
        structured_data = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        print("Failed to decode JSON:", e)
        return None
    print(structured_data)
    return structured_data


def create_docx(data_list):
    # Define the order of major heads
    branch_order = {
        'E': 'Executive Branch',
        'L': 'Legislative Branch',
        'S': 'State Officials',
        'O': 'Others'
    }

    # Sort data_list based on branch_head order
    data_list.sort(key=lambda x: list(branch_order.keys()).index(x['branch_head']))

    doc = Document()
    
    # Add header "DEA Quotes" with increased font size and centered
    header = doc.add_heading('DEA Quotes', level=3)
    run_header = header.runs[0]
    run_header.font.size = Pt(40)
    run_header.font.name = 'Times New Roman'  # Set Times New Roman as default font
    run_header._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run_header.font.color.rgb = RGBColor(128, 128, 128) 
    header.alignment = 1  # Center alignment
    

    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = 1  # Center alignment
    date_run = date_paragraph.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(14)  # Adjust font size if needed
    # date_run.font.color.rgb = RGBColor(128, 128, 128) 

    for branch, branch_name in branch_order.items():
        # Filter items by branch_head
        branch_items = [item for item in data_list if item['branch_head'] == branch]

        if branch_items:
            # Add branch name as a heading
            branch_heading = doc.add_heading(level=2)
            run_branch_heading = branch_heading.add_run(branch_name)
            run_branch_heading.bold = True
            run_branch_heading.font.size = Pt(16)
            run_branch_heading.font.name = 'Times New Roman'
            run_branch_heading.font.color.rgb = RGBColor(0, 0, 0)  # Black color

            for item in branch_items:
                # Add the topic to the document
                topic = item['topic']
                p_topic = doc.add_paragraph()
                run_topic_label = p_topic.add_run("Topic: ")
                run_topic_label.bold = True
                run_topic_label.font.name = 'Times New Roman'
                run_topic_content = p_topic.add_run(topic)
                run_topic_content.font.name = 'Times New Roman'

                # Add quoted information
                quoted = item.get('quoted', 'No quotes available')
                p_quoted = doc.add_paragraph()
                run_quoted_label = p_quoted.add_run("Quoted: ")
                run_quoted_label.bold = True
                run_quoted_label.font.name = 'Times New Roman'
                run_quoted_content = p_quoted.add_run(quoted)
                run_quoted_content.font.name = 'Times New Roman'

                # Add background information
                background = item.get('background', 'No background available')
                p_background = doc.add_paragraph()
                run_background_label = p_background.add_run("Background: ")
                run_background_label.bold = True
                run_background_label.font.name = 'Times New Roman'
                run_background_content = p_background.add_run(background)
                run_background_content.font.name = 'Times New Roman'

                # Add names and their quotes with extra spacing
                people_quotes = item.get('people_quotes', [])
                if not people_quotes:
                    p_no_quotes = doc.add_paragraph()
                    p_no_quotes.paragraph_format.space_after = Pt(0) 
                    run_no_quotes = p_no_quotes.add_run("No quotes found in people_quotes.")
                    run_no_quotes.font.name = 'Times New Roman'
                else:
                    for entry in people_quotes:
                        # Add extra space before name
                        doc.add_paragraph()
                        name = entry.get('name', 'Unknown')
                        p_name = doc.add_paragraph()
                        p_name.paragraph_format.space_after = Pt(0)
                        run_name = p_name.add_run(name)
                        run_name.bold = True
                        run_name.font.name = 'Times New Roman'

                        # Add quote with proper formatting
                        p_quote = doc.add_paragraph()
                        p_quote.paragraph_format.space_after = Pt(0) 
                        run_quote = p_quote.add_run("Quote: ")  # Added colon after Quote
                        run_quote.bold = True
                        run_quote.font.size = Pt(14)  # Set to 14pt
                        run_quote.font.name = 'Times New Roman'
                        
                        quote = entry.get('quote', 'No quote available')
                        run_quotes = p_quote.add_run(quote)
                        # run_quotes.paragraph_format.space_after = Pt(0) 
                        run_quotes.italic = True
                        run_quotes.font.name = 'Times New Roman'

                # Add date with tighter spacing
                date = item.get('date', 'No date available')
                p_date = doc.add_paragraph()
                p_date.paragraph_format.space_before = Pt(0)  # No space before
                p_date.paragraph_format.space_after = Pt(0) # Reduce space before date
                run_date = p_date.add_run(date)
                run_date.font.name = 'Times New Roman Italic'

                # Add link
                link = item.get('link', 'No link available')
                p_link = doc.add_paragraph()
                p_link.paragraph_format.space_before = Pt(0) 
                p_link.paragraph_format.space_after = Pt(6)# Reduce space before link
                run_link = p_link.add_run(link)
                run_link.font.color.rgb = RGBColor(0, 0, 255)
                run_link.font.underline = True
                run_link.font.name = 'Times New Roman Italic'

                # Add separator with consistent spacing
                p_separator = doc.add_paragraph("---")
                p_separator.paragraph_format.space_before = Pt(6)  # Small space before
                p_separator.paragraph_format.space_after = Pt(12)
                for run in p_separator.runs:
                    run.font.name = 'Times New Roman'

    # Save the document to a file
    doc_path = "newsletter_output.docx"
    doc.save(doc_path)

    return doc_path


# def extract_image_from_page(pdf_document, page_num):
#     """Extract image from a single PDF page."""
#     try:
#         page = pdf_document[page_num]
#         pix = page.get_pixmap(matrix=fitz.Matrix(100/72, 100/72))  # Reduce resolution
#         img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#         img = img.resize((int(pix.width * 0.5), int(pix.height * 0.5)))  # Resize to 50%
#         return np.array(img)
#     except Exception as e:
#         print(f"Error extracting image from page {page_num}: {e}")
#         return None

# def extract_images_from_pdf(pdf_file):
#     """Extract images from all pages of the PDF."""
#     images = []
#     try:
#         pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
#         with ThreadPoolExecutor() as executor:
#             futures = [
#                 executor.submit(extract_image_from_page, pdf_document, page_num) 
#                 for page_num in range(len(pdf_document))
#             ]
#             for future in futures:
#                 result = future.result()
#                 if result is not None:
#                     images.append(result)
#         return images
#     except Exception as e:
#         print(f"Error extracting images from PDF: {e}")
#         return []


@st.cache_resource(show_spinner=False)
def load_easyocr():
    """Initialize EasyOCR reader with caching"""
    return easyocr.Reader(['en'], gpu=False)

def process_ocr_pdf(pdf_file):
    """Optimized function to process PDF and extract text using OCR."""
    try:
        reader = load_easyocr()
        pdf_file.seek(0)
        
        # Open PDF
        pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
        all_text = []
        
        # Process all pages
        for page_num in range(len(pdf_document)):
            try:
                page = pdf_document[page_num]
                # Use consistent 150 DPI for all pages
                matrix = fitz.Matrix(200/72, 200/72)  
                # Get page image
                pix = page.get_pixmap(matrix=matrix)
                # print(f"Page {page_num} - Image size: {pix.width}x{pix.height}")  # Debug size
                # Convert to PIL Image
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                # Optional: Enhance contrast (helps EasyOCR with faint text)
                # img = img.convert("L").convert("RGB")  # Convert to grayscale and back to RGB
                # Convert to numpy array
                img_np = np.array(img)
                # Extract text
                results = reader.readtext(img_np,paragraph=True, low_text=0.5)
                # print({"page": page_num, "results": results})
                # Extract and clean text
                page_text = " ".join([entry[1] for entry in results if entry[1].strip()])
                # print({"page": page_num, "page_text": page_text})
                # Add to collection
                if page_text.strip():
                    all_text.append(page_text)
            except Exception as e:
                print(f"Error processing page {page_num}: {e}")
                continue
        
        # Combine all text
        final_text_1 = " ".join(all_text).strip()
        # print({"final_text_1": final_text_1})
        # print({"all_text": all_text})
        
        # Basic text cleaning
        final_text_2 = re.sub(r'\s+', ' ', final_text_1)  # Remove multiple spaces
        final_text = re.sub(r'[^\x00-\x7F]+', '', final_text_2)  # Remove non-ASCII characters
        # print({"final_text": final_text})
        return final_text if final_text else None

    except Exception as e:
        print(f"Failed to process the file: {e}")
        return None